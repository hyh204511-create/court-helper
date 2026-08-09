"""Build polished DOCX handoff manuals from docs/delivery Markdown sources."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "delivery"
OUTPUT_DIR = ROOT / "output" / "docs"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
FONT = "Microsoft YaHei"
MONO = "Consolas"


def set_run_font(run, *, size=10.5, bold=False, color="1F2937", name=FONT, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run._r.extend([begin, instruction, separate, display, end])
    set_run_font(field_run, size=8.5, color=MID_GRAY)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=8.5, color=MID_GRAY)


def configure_document(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("court-helper  |  客户交付资料")
    set_run_font(run, size=8.5, bold=True, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    add_page_field(footer)

    doc.core_properties.title = running_title
    doc.core_properties.subject = "法院立案/强执查询助手客户交付文档"
    doc.core_properties.author = "court-helper project"
    doc.core_properties.last_modified_by = "court-helper project"


def inline_runs(paragraph, text, *, base_size=10.5, color="1F2937"):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=9.5, name=MONO, color=NAVY)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, color=color)


def add_cover(doc, product, title, metadata):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("客户交付手册")
    set_run_font(run, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(product)
    set_run_font(run, size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(title)
    set_run_font(run, size=15, bold=True, color=BLUE)

    for line in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        inline_runs(p, line, base_size=9.5, color=MID_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-+:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    usable = 9360
    widths = [usable // cols] * cols
    widths[-1] += usable - sum(widths)
    for row_index, values in enumerate(rows):
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            inline_runs(p, values[col_index] if col_index < len(values) else "", base_size=8.8)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    inline_runs(p, text, base_size=9.5, color=NAVY)
    set_table_geometry(table, [9360])


def build_document(source):
    lines = source.read_text(encoding="utf-8").splitlines()
    headings = [(i, line.lstrip("#").strip()) for i, line in enumerate(lines) if line.startswith("#")]
    product = headings[0][1]
    title = headings[1][1] if len(headings) > 1 else product
    metadata = []
    cursor = headings[1][0] + 1 if len(headings) > 1 else headings[0][0] + 1
    while cursor < len(lines) and not lines[cursor].startswith("## "):
        if lines[cursor].strip() and not lines[cursor].lstrip().startswith(">"):
            metadata.append(lines[cursor].strip().rstrip("  "))
        cursor += 1

    doc = Document()
    configure_document(doc, title)
    add_cover(doc, product, title, metadata)

    in_code = False
    code_lines = []
    index = cursor
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.right_indent = Inches(0.2)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, size=8.5, name=MONO, color=NAVY)
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT_GRAY)
                p_pr.append(shd)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if stripped.startswith(">"):
            add_callout(doc, stripped.lstrip("> "))
        elif raw.startswith("### "):
            doc.add_paragraph(raw[4:].strip(), style="Heading 3")
        elif raw.startswith("## "):
            doc.add_paragraph(raw[3:].strip(), style="Heading 1")
        elif raw.startswith("# "):
            pass
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            inline_runs(p, re.sub(r"^\d+\.\s*", "", stripped))
        elif stripped.startswith("- [ ] "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            inline_runs(p, "☐ " + stripped[6:])
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            inline_runs(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            inline_runs(p, stripped)
        index += 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / f"{source.stem}.docx"
    doc.save(output)
    return output


def main():
    sources = sorted(path for path in SOURCE_DIR.glob("*.md") if path.name != "README.md")
    if not sources:
        raise SystemExit("No delivery Markdown sources found")
    outputs = [build_document(source) for source in sources]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    sys.exit(main())
